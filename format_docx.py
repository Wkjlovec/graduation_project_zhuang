import re
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT = "docs/论文修改.docx"
OUTPUT = "docs/论文修改.docx"

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TNR = "Times New Roman"

PT_ERHAO = Pt(22)
PT_XIAOER = Pt(18)
PT_XIAOSAN = Pt(15)
PT_SIHAO = Pt(14)
PT_XIAOSI = Pt(12)
PT_WUHAO = Pt(10.5)

CHAPTER_TITLES = {
    "第一章 绪论", "第二章 相关技术及方法", "第三章 系统需求分析",
    "第四章 系统详细设计与实现", "第五章 系统测试", "第六章 总结与展望",
}
SECTION_RE = re.compile(r"^([1-6])\.([0-9]+)\s+\S")


def is_section_heading(text):
    if not SECTION_RE.match(text):
        return False
    words = text.split(None, 1)
    if len(words) < 2:
        return False
    rest = words[1]
    if rest[0].isdigit():
        return False
    skip_kw = ("梳理", "描述", "列出", "按功能", "指出")
    if "节" in rest[:4] and any(k in rest for k in skip_kw):
        return False
    return True


def set_run_font(run, cn, en, size, bold=False):
    run.font.size = size
    run.font.bold = bold
    run.font.name = en
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)
    rFonts.set(qn("w:eastAsia"), cn)


def set_pf(para, align, line_spacing=1.5, indent_cm=None,
           space_before_pt=0, space_after_pt=0):
    pf = para.paragraph_format
    pf.alignment = align
    pf.line_spacing = line_spacing
    pf.first_line_indent = Cm(indent_cm) if indent_cm else None
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)


def fmt_runs(para, cn, en, size, bold=False):
    for run in para.runs:
        set_run_font(run, cn, en, size, bold)


def ensure_runs(para, cn, en, size, bold=False):
    if not para.runs:
        para.add_run(para.text)
    fmt_runs(para, cn, en, size, bold)


def set_outline_level(para, level):
    """Set outline level (0=Heading1, 1=Heading2, etc.) so TOC field can find it."""
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._element.insert(0, pPr)
    outlineLvl = pPr.find(qn("w:outlineLvl"))
    if outlineLvl is None:
        outlineLvl = OxmlElement("w:outlineLvl")
        pPr.append(outlineLvl)
    outlineLvl.set(qn("w:val"), str(level))


def is_figure_line(t):
    return bool(re.match(r"^图\s*[0-9]+-[0-9]+", t))


def is_table_caption(t):
    return bool(re.match(r"^表\s*[0-9]+-[0-9]+", t))


def is_ref_entry(t):
    return bool(re.match(r"^\[[0-9]+\]", t))


def make_toc_field_para():
    """Create a paragraph containing a TOC field code."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "312")   # 1.25 * 240 = 300, use 312 for ~1.3x
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    p.append(pPr)

    def make_run():
        return OxmlElement("w:r")

    r1 = make_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r1.append(fc1)
    p.append(r1)

    r2 = make_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    r2.append(instr)
    p.append(r2)

    r3 = make_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    r3.append(fc2)
    p.append(r3)

    r4 = make_run()
    rPr4 = OxmlElement("w:rPr")
    sz4 = OxmlElement("w:sz")
    sz4.set(qn("w:val"), "24")
    rPr4.append(sz4)
    szCs4 = OxmlElement("w:szCs")
    szCs4.set(qn("w:val"), "24")
    rPr4.append(szCs4)
    rFonts4 = OxmlElement("w:rFonts")
    rFonts4.set(qn("w:eastAsia"), FONT_SONG)
    rFonts4.set(qn("w:ascii"), FONT_TNR)
    rFonts4.set(qn("w:hAnsi"), FONT_TNR)
    rPr4.append(rFonts4)
    r4.append(rPr4)
    t4 = OxmlElement("w:t")
    t4.text = "[TOC placeholder - right-click and Update Field in Word]"
    r4.append(t4)
    p.append(r4)

    r5 = make_run()
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    r5.append(fc3)
    p.append(r5)

    return p


def make_page_break():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


# ── Load & Format ──
doc = Document(INPUT)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    if i == 0 and "基于微服务架构" in text:
        set_pf(para, WD_ALIGN_PARAGRAPH.CENTER,
               space_before_pt=22, space_after_pt=44)
        ensure_runs(para, FONT_HEI, FONT_TNR, PT_ERHAO, bold=True)
        continue

    if text == "摘 要":
        set_pf(para, WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=22)
        ensure_runs(para, FONT_HEI, FONT_TNR, PT_XIAOER, bold=True)
        set_outline_level(para, 0)
        continue

    if text.startswith("Design and Implementation"):
        set_pf(para, WD_ALIGN_PARAGRAPH.CENTER,
               space_before_pt=22, space_after_pt=44)
        ensure_runs(para, FONT_TNR, FONT_TNR, PT_ERHAO, bold=True)
        continue

    if text == "ABSTRACT":
        set_pf(para, WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=22)
        ensure_runs(para, FONT_TNR, FONT_TNR, PT_XIAOER, bold=True)
        set_outline_level(para, 0)
        continue

    if text.startswith("关键词") or text.startswith("**关键词"):
        set_pf(para, WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=0.85)
        for run in para.runs:
            if "关键词" in run.text:
                set_run_font(run, FONT_HEI, FONT_TNR, PT_XIAOSI, bold=True)
            else:
                set_run_font(run, FONT_SONG, FONT_TNR, PT_XIAOSI)
        continue

    if text.startswith("Key Words") or text.startswith("**Key Words"):
        set_pf(para, WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=1.69)
        fmt_runs(para, FONT_TNR, FONT_TNR, PT_XIAOSI)
        continue

    if text in CHAPTER_TITLES:
        set_pf(para, WD_ALIGN_PARAGRAPH.CENTER)
        ensure_runs(para, FONT_HEI, FONT_TNR, PT_XIAOER, bold=True)
        set_outline_level(para, 0)
        continue

    if text == "参 考 文 献":
        set_pf(para, WD_ALIGN_PARAGRAPH.CENTER)
        ensure_runs(para, FONT_HEI, FONT_TNR, PT_XIAOER, bold=True)
        set_outline_level(para, 0)
        continue

    if text in ("致谢", "致 谢"):
        para.clear()
        run = para.add_run("致 谢")
        set_pf(para, WD_ALIGN_PARAGRAPH.CENTER)
        set_run_font(run, FONT_HEI, FONT_TNR, PT_XIAOER, bold=True)
        set_outline_level(para, 0)
        continue

    if is_section_heading(text):
        set_pf(para, WD_ALIGN_PARAGRAPH.LEFT)
        ensure_runs(para, FONT_HEI, FONT_TNR, PT_XIAOSAN, bold=True)
        set_outline_level(para, 1)
        continue

    if is_ref_entry(text):
        set_pf(para, WD_ALIGN_PARAGRAPH.JUSTIFY)
        fmt_runs(para, FONT_SONG, FONT_TNR, PT_WUHAO)
        continue

    if is_figure_line(text):
        set_pf(para, WD_ALIGN_PARAGRAPH.CENTER)
        ensure_runs(para, FONT_SONG, FONT_TNR, PT_WUHAO)
        continue

    if is_table_caption(text):
        set_pf(para, WD_ALIGN_PARAGRAPH.CENTER)
        ensure_runs(para, FONT_SONG, FONT_TNR, PT_WUHAO)
        continue

    if not text:
        continue

    # Chinese abstract
    if i == 2:
        set_pf(para, WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=0.85)
        fmt_runs(para, FONT_SONG, FONT_TNR, PT_XIAOSI)
        continue

    # English abstract
    if i == 6:
        set_pf(para, WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=1.69)
        fmt_runs(para, FONT_TNR, FONT_TNR, PT_XIAOSI)
        continue

    # Default: body text
    set_pf(para, WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=0.85)
    fmt_runs(para, FONT_SONG, FONT_TNR, PT_XIAOSI)

# Format tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                set_pf(para, WD_ALIGN_PARAGRAPH.CENTER, indent_cm=None)
                fmt_runs(para, FONT_SONG, FONT_TNR, PT_WUHAO)

# ── Insert TOC after Key Words ──
kw_elem = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith("Key Words"):
        kw_elem = para._element
        break

if kw_elem is not None:
    pb1 = make_page_break()
    kw_elem.addnext(pb1)

    toc_title_p = OxmlElement("w:p")
    toc_pPr = OxmlElement("w:pPr")
    toc_jc = OxmlElement("w:jc")
    toc_jc.set(qn("w:val"), "center")
    toc_pPr.append(toc_jc)
    toc_sp = OxmlElement("w:spacing")
    toc_sp.set(qn("w:after"), "440")  # ~22pt in twips (22*20=440)
    toc_sp.set(qn("w:line"), "360")   # 1.5x line spacing
    toc_sp.set(qn("w:lineRule"), "auto")
    toc_pPr.append(toc_sp)
    toc_title_p.append(toc_pPr)

    toc_r = OxmlElement("w:r")
    toc_rPr = OxmlElement("w:rPr")
    toc_b = OxmlElement("w:b")
    toc_rPr.append(toc_b)
    toc_sz = OxmlElement("w:sz")
    toc_sz.set(qn("w:val"), "36")  # 18pt = 小二号
    toc_rPr.append(toc_sz)
    toc_szCs = OxmlElement("w:szCs")
    toc_szCs.set(qn("w:val"), "36")
    toc_rPr.append(toc_szCs)
    toc_rF = OxmlElement("w:rFonts")
    toc_rF.set(qn("w:eastAsia"), FONT_HEI)
    toc_rF.set(qn("w:ascii"), FONT_TNR)
    toc_rF.set(qn("w:hAnsi"), FONT_TNR)
    toc_rPr.append(toc_rF)
    toc_r.append(toc_rPr)
    toc_t = OxmlElement("w:t")
    toc_t.text = "目 录"
    toc_r.append(toc_t)
    toc_title_p.append(toc_r)
    pb1.addnext(toc_title_p)

    toc_field = make_toc_field_para()
    toc_title_p.addnext(toc_field)

    pb2 = make_page_break()
    toc_field.addnext(pb2)

doc.save(OUTPUT)
print(f"Done: {OUTPUT}")
print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
