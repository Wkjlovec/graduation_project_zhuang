# -*- coding: utf-8 -*-
"""Add argumentative commentary paragraphs to the docx."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_SONG = "\u5b8b\u4f53"
FONT_TNR = "Times New Roman"
PT_XIAOSI = Pt(12)


def set_run_font(run, cn, en, size):
    run.font.size = size
    run.font.name = en
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)
    rFonts.set(qn("w:eastAsia"), cn)


def make_body_para(text):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    pPr.append(jc)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "482")
    pPr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "24")
    rPr.append(szCs)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_SONG)
    rFonts.set(qn("w:ascii"), FONT_TNR)
    rFonts.set(qn("w:hAnsi"), FONT_TNR)
    rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


doc = Document("docs/\u8bba\u6587\u4fee\u6539.docx")

# === Commentary 1: Insert after 1.1's "短板" paragraph, before "面向上述问题" ===
# Current para 14 contains ALL of 1.1 in one paragraph. We need to split it
# and insert commentary in the middle.

p14 = doc.paragraphs[14]
full_text = p14.text

# Split point: after "交付质量缺少可量化的判定依据。" and before "面向上述问题"
split_marker = "\u4ea4\u4ed8\u8d28\u91cf\u7f3a\u5c11\u53ef\u91cf\u5316\u7684\u5224\u5b9a\u4f9d\u636e\u3002"
idx = full_text.find(split_marker)
if idx < 0:
    print("ERROR: split marker not found in para 14")
else:
    part1 = full_text[:idx + len(split_marker)]
    part2 = full_text[idx + len(split_marker):].strip()

    # Replace para 14 with part1
    p14.clear()
    run = p14.add_run(part1)
    set_run_font(run, FONT_SONG, FONT_TNR, PT_XIAOSI)
    p14.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p14.paragraph_format.first_line_indent = Cm(0.85)
    p14.paragraph_format.line_spacing = 1.5

    commentary1 = (
        "\u672c\u6587\u6ce8\u610f\u5230\uff0c\u76ee\u524d\u9ad8\u6821\u573a\u666f\u4e0b"
        "\u5bf9\u4e0a\u8ff0\u95ee\u9898\u7684\u5e38\u89c1\u5e94\u5bf9\u65b9\u5f0f"
        "\u5e76\u4e0d\u4ee4\u4eba\u6ee1\u610f\u3002"
        "\u4e00\u79cd\u505a\u6cd5\u662f\u201c\u5c01\u88c5\u800c\u4e0d\u62c6\u5206\u201d"
        "\u2014\u2014\u5c06\u6240\u6709\u529f\u80fd\u5806\u780c\u5728\u540c\u4e00\u4e2a"
        "\u5de5\u7a0b\u4e2d\uff0c\u7528\u66f4\u9ad8\u914d\u7f6e\u7684\u670d\u52a1\u5668"
        "\u6765\u786c\u6297\u5e76\u53d1\u538b\u529b\u3002"
        "\u8fd9\u4e0e\u667a\u80fd\u624b\u673a\u884c\u4e1a\u4ee5\u589e\u5927\u5185\u5b58"
        "\u6765\u5bb9\u5fcd\u540e\u53f0\u5e38\u9a7b\u5e94\u7528\u7684\u601d\u8def\u5982\u51fa\u4e00\u8f99"
        "\u2014\u2014\u786c\u4ef6\u6027\u80fd\u7684\u63d0\u5347\u786e\u5b9e\u5ef6\u7f13\u4e86\u77db\u76fe\uff0c"
        "\u4f46\u5e76\u672a\u89e3\u51b3\u67b6\u6784\u8026\u5408\u7684\u6839\u672c\u95ee\u9898\uff0c"
        "\u53cd\u800c\u4f7f\u5f97\u7cfb\u7edf\u7684\u8fd0\u7ef4\u6210\u672c\u548c\u8fed\u4ee3\u96be\u5ea6"
        "\u968f\u529f\u80fd\u589e\u957f\u800c\u6301\u7eed\u7d2f\u79ef\u3002"
        "\u53e6\u4e00\u79cd\u505a\u6cd5\u662f\u201c\u53ea\u505a\u524d\u7aef\u4e0d\u52a8\u540e\u7aef\u201d"
        "\u2014\u2014\u5728\u539f\u6709\u5355\u4f53\u5de5\u7a0b\u4e4b\u4e0a\u589e\u52a0\u4e00\u5957\u79fb\u52a8\u7aef\u9875\u9762\uff0c"
        "\u7528 CSS \u5a92\u4f53\u67e5\u8be2\u505a\u54cd\u5e94\u5f0f\u9002\u914d\u3002"
        "\u8fd9\u79cd\u65b9\u5f0f\u5728\u89c6\u89c9\u4e0a\u89e3\u51b3\u4e86\u53cc\u7aef\u95ee\u9898\uff0c"
        "\u4f46\u524d\u540e\u7aef\u4ecd\u7136\u8026\u5408\u5728\u4e00\u8d77\uff0c"
        "\u4e00\u6b21\u5c40\u90e8\u4fee\u6539\u4ecd\u7136\u9700\u8981\u5168\u5c40\u91cd\u65b0\u90e8\u7f72\uff0c"
        "\u8d28\u91cf\u4fdd\u969c\u548c\u72ec\u7acb\u8fed\u4ee3\u7684\u95ee\u9898\u539f\u5c01\u4e0d\u52a8\u3002"
        "\u8fd9\u4e24\u79cd\u5e94\u5bf9\u65b9\u5f0f\u90fd\u662f\u5728\u56de\u907f\u67b6\u6784\u5c42\u9762\u7684\u6839\u672c\u6539\u9020\uff0c"
        "\u5b83\u4eec\u80fd\u591f\u5ef6\u7f13\u95ee\u9898\u7684\u7206\u53d1\uff0c"
        "\u5374\u65e0\u6cd5\u6d88\u9664\u95ee\u9898\u672c\u8eab\u3002"
    )

    # Insert commentary paragraph after para 14
    new_p1 = make_body_para(commentary1)
    p14._element.addnext(new_p1)

    # Insert part2 (the "面向上述问题..." paragraph) after commentary
    new_p2 = make_body_para(part2)
    new_p1.addnext(new_p2)

    print(f"1.1: Split para 14 and inserted commentary ({len(commentary1)} chars)")

# === Commentary 2: Insert in 1.2 after the decline paragraph ===
# Find the paragraph about forum decline (starts with "2009")
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("2009") and "\u5929\u6daf\u793e\u533a" in p.text:
        decline_para = p

        commentary2 = (
            "\u8bba\u575b\u7684\u8870\u9000\u5f15\u53d1\u4e86\u4e00\u4e2a\u503c\u5f97\u601d\u8003\u7684\u95ee\u9898\uff1a"
            "\u201c\u4ee5\u4e3b\u9898\u805a\u5408\u8ba8\u8bba\u201d\u8fd9\u4e00\u9700\u6c42\u7a76\u7adf\u662f\u88ab\u65b0\u5e73\u53f0\u66ff\u4ee3\u4e86\uff0c"
            "\u8fd8\u662f\u88ab\u65e7\u6280\u672f\u62d6\u7d2f\u4e86\uff1f"
            "\u4ece Reddit \u548c Stack Overflow \u7684\u6301\u7eed\u6d3b\u8dc3\u6765\u770b\uff0c"
            "\u7b54\u6848\u663e\u7136\u662f\u540e\u8005\u3002"
            "\u7528\u6237\u5e76\u6ca1\u6709\u653e\u5f03\u5bf9\u7ed3\u6784\u5316\u8ba8\u8bba\u548c\u77e5\u8bc6\u6c89\u6dc0\u7684\u9700\u6c42\uff0c"
            "\u4ed6\u4eec\u653e\u5f03\u7684\u662f\u754c\u9762\u9648\u65e7\u3001\u4e0d\u9002\u914d\u79fb\u52a8\u7aef\u3001"
            "\u7f3a\u4e4f\u73b0\u4ee3\u4ea4\u4e92\u7279\u6027\u7684\u8001\u65e7\u8bba\u575b\u8f6f\u4ef6\u3002"
            "\u6362\u8a00\u4e4b\uff0c\u95ee\u9898\u4e0d\u5728\u4e8e\u8bba\u575b\u8fd9\u4e00\u4ea7\u54c1\u5f62\u6001\u672c\u8eab\uff0c"
            "\u800c\u5728\u4e8e\u627f\u8f7d\u5b83\u7684\u6280\u672f\u67b6\u6784\u672a\u80fd\u8ddf\u4e0a\u65f6\u4ee3\u3002"
            "\u8fd9\u6070\u6070\u8bf4\u660e\uff0c\u5728\u9ad8\u6821\u573a\u666f\u4e2d\u91cd\u65b0\u5ba1\u89c6\u8bba\u575b\u7cfb\u7edf\u7684\u6280\u672f\u67b6\u6784\uff0c"
            "\u7528\u5fae\u670d\u52a1\u3001\u524d\u540e\u7aef\u5206\u79bb\u548c\u79fb\u52a8\u7aef\u9002\u914d\u7b49\u73b0\u4ee3\u5de5\u7a0b\u5b9e\u8df5"
            "\u6765\u91cd\u65b0\u6fc0\u6d3b\u8fd9\u4e00\u4ea7\u54c1\u5f62\u6001\uff0c"
            "\u5177\u6709\u5b9e\u9645\u7684\u5de5\u7a0b\u4ef7\u503c\u548c\u7814\u7a76\u610f\u4e49\u3002"
        )

        new_p = make_body_para(commentary2)
        decline_para._element.addnext(new_p)
        print(f"1.2: Inserted commentary after decline paragraph ({len(commentary2)} chars)")
        break

doc.save("docs/\u8bba\u6587\u4fee\u6539.docx")
print("\nDone.")
