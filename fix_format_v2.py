"""
Fix formatting differences based on the updated standard:
1. Keywords: separator ; not , ; on new line, left-aligned, indent 2 chars
2. H1 (chapter titles): 小二号黑体居中, 段前段后空1行, 标号后空1格
3. H2 (section headings): 四号黑体(14pt) not 小三号(15pt), 段前空1行, 标号后空2格
4. H3 (subsection): 小四号黑体(12pt), 段前空1行, 标号后空2格
5. 参考文献/致谢 special headings: same as H1 spacing
"""
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TNR = "Times New Roman"

PT_XIAOER = Pt(18)
PT_SIHAO = Pt(14)
PT_XIAOSI = Pt(12)
PT_WUHAO = Pt(10.5)

ONE_LINE_BEFORE = Pt(22)
ONE_LINE_AFTER = Pt(22)

doc = Document("docs/论文修改.docx")

CHAPTER_TITLES = {
    "第一章 绪论", "第二章 相关技术及方法", "第三章 系统需求分析",
    "第四章 系统详细设计与实现", "第五章 系统测试", "第六章 总结与展望",
}

SECTION_RE = re.compile(r"^([1-6])\.([0-9]+)\s+\S")


def is_section_heading(text):
    if not SECTION_RE.match(text):
        return False
    words = text.split(None, 1)
    if len(words) < 2:
        return False
    rest = words[1]
    if rest[0].isdigit():
        return False
    skip_kw = ("梳理", "描述", "列出", "按功能", "指出")
    if "节" in rest[:4] and any(k in rest for k in skip_kw):
        return False
    return True


def set_run_font(run, cn, en, size, bold=False):
    run.font.size = size
    run.font.bold = bold
    run.font.name = en
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)
    rFonts.set(qn("w:eastAsia"), cn)


fixes = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    pf = para.paragraph_format

    # === Fix 1: Chinese keywords - separator ; and left-align ===
    if text.startswith("关键词") or text.startswith("**关键词"):
        old_text = para.text
        new_text = old_text.replace("，", "；")
        if "，" in old_text:
            para.clear()
            run_label = para.add_run("关键词：")
            set_run_font(run_label, FONT_HEI, FONT_TNR, PT_XIAOSI, bold=True)
            kw_content = new_text.split("：", 1)[1].strip() if "：" in new_text else new_text
            kw_content = kw_content.replace("**", "").strip()
            run_content = para.add_run(kw_content)
            set_run_font(run_content, FONT_SONG, FONT_TNR, PT_XIAOSI)
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0.85)
            fixes.append(f"Fix KW-CN at {i}: separator , -> ;, left-aligned")
        continue

    # === Fix 1b: English keywords - separator ; ===
    if text.startswith("Key Words") or text.startswith("**Key Words"):
        old_text = para.text
        if ", " in old_text:
            new_text = old_text.replace(", ", "; ")
            para.clear()
            run = para.add_run(new_text.replace("**", ""))
            set_run_font(run, FONT_TNR, FONT_TNR, PT_XIAOSI)
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(1.69)
            fixes.append(f"Fix KW-EN at {i}: separator , -> ;")
        continue

    # === Fix 2: Chapter titles (H1) - 段前段后空1行 ===
    if text in CHAPTER_TITLES:
        pf.space_before = ONE_LINE_BEFORE
        pf.space_after = ONE_LINE_AFTER
        fixes.append(f"Fix H1 at {i}: added space_before/after 22pt \"{text}\"")
        continue

    # === Fix 2b: 参考文献/致谢 - same spacing as H1 ===
    if text == "参 考 文 献" or text == "致 谢":
        pf.space_before = ONE_LINE_BEFORE
        pf.space_after = ONE_LINE_AFTER
        fixes.append(f"Fix special H1 at {i}: added space_before/after \"{text}\"")
        continue

    # === Fix 3: Section headings (H2) - 四号(14pt), 段前空1行 ===
    if is_section_heading(text):
        for run in para.runs:
            run.font.size = PT_SIHAO
        pf.space_before = ONE_LINE_BEFORE
        pf.space_after = Pt(0)
        fixes.append(f"Fix H2 at {i}: 15pt->14pt, space_before=22pt \"{text[:30]}\"")
        continue

doc.save("docs/论文修改.docx")
print(f"Applied {len(fixes)} fixes:")
for f in fixes:
    print(f"  {f}")
