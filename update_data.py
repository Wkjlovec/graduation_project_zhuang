from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_SONG = "宋体"
FONT_TNR = "Times New Roman"
PT_WUHAO = Pt(10.5)

def set_run_font(run, cn, en, size):
    run.font.size = size
    run.font.bold = False
    run.font.name = en
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)
    rFonts.set(qn("w:eastAsia"), cn)

def set_cell_text(cell, text):
    for p in cell.paragraphs:
        p.clear()
        run = p.add_run(text)
        set_run_font(run, FONT_SONG, FONT_TNR, PT_WUHAO)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_row(table, values):
    row = table.add_row()
    for ci, val in enumerate(values):
        set_cell_text(row.cells[ci], str(val))
    return row

doc = Document("docs/论文修改.docx")

# === Table 5-3: Cache benchmark ===
t53 = doc.tables[11]
cache_data = [
    "59.12 ms",
    "14.98 ms",
    "4.33 ms",
    "27.81 ms",
    "26.3 ms",
    "74.66%",
]
for ri, val in enumerate(cache_data):
    set_cell_text(t53.rows[ri + 1].cells[1], val)

# === Table 5-4: Concurrent load ===
t54 = doc.tables[12]

row1_data = ["200", "0", "100.0%", "8.02", "10.37", "1186.17"]
row2_data = ["1000", "0", "100.0%", "33.12", "50.62", "1420.12"]
row3_data = ["1999", "1", "99.95%", "59.0", "97.44", "1391.7"]

for ci, val in enumerate(row1_data):
    set_cell_text(t54.rows[1].cells[ci + 2], val)
for ci, val in enumerate(row2_data):
    set_cell_text(t54.rows[2].cells[ci + 2], val)
for ci, val in enumerate(row3_data):
    set_cell_text(t54.rows[3].cells[ci + 2], val)

add_row(t54, ["500", "10000", "10000", "0", "100.0%", "124.35", "226.01", "1556.64"])
add_row(t54, ["1000", "20000", "19998", "2", "99.99%", "191.35", "365.98", "1450.42"])

# === Update body text: concurrent test description ===
for i, para in enumerate(doc.paragraphs):
    t = para.text
    if "测试分三轮进行，并发用户数依次为 10、50 和 100" in t:
        old = "测试分三轮进行，并发用户数依次为 10、50 和 100，每个虚拟用户向帖子列表接口发送 20 次 GET 请求"
        new = "测试分五轮进行，并发用户数依次为 10、50、100、500 和 1000，每个虚拟用户向帖子列表接口发送 20 次 GET 请求"
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                break
        else:
            full = para.text
            if old in full:
                para.clear()
                run = para.add_run(full.replace(old, new))
                set_run_font(run, FONT_SONG, FONT_TNR, Pt(12))
        print(f"  Updated concurrent description at para {i}")
        break

# Update "10 至 100 并发" mentions in abstract, chapter 5.4, chapter 6.1
updates = [
    ("10 至 100 并发用户下验证了系统的承载能力", "10 至 1000 并发用户下验证了系统的承载能力"),
    ("10 至 100 并发用户的规模下验证了系统能够维持稳定的响应", "10 至 1000 并发用户的规模下验证了系统能够维持稳定的响应"),
    ("系统在 10 至 100 并发用户下的承载能力", "系统在 10 至 1000 并发用户下的承载能力"),
    ("10 to 100 users verify", "10 to 1000 users verify"),
    ("under 10 to 100 users", "under 10 to 1000 users"),
]

for old_text, new_text in updates:
    for i, para in enumerate(doc.paragraphs):
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    print(f"  Updated '{old_text[:30]}...' at para {i}")
                    break
            else:
                full = para.text
                if old_text in full:
                    new_full = full.replace(old_text, new_text)
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = new_full
                    else:
                        para.add_run(new_full)
                    print(f"  Updated (full replace) '{old_text[:30]}...' at para {i}")

doc.save("docs/论文修改.docx")
print("\nDone!")
