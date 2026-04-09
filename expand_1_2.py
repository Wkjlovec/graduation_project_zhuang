# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_SONG = "宋体"
FONT_TNR = "Times New Roman"
PT_XIAOSI = Pt(12)


def set_run_font(run, cn, en, size, bold=False):
    run.font.size = size
    run.font.bold = bold
    run.font.name = en
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)
    rFonts.set(qn("w:eastAsia"), cn)


def make_body_para(text):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    pPr.append(jc)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "482")
    pPr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "24")
    rPr.append(szCs)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_SONG)
    rFonts.set(qn("w:ascii"), FONT_TNR)
    rFonts.set(qn("w:hAnsi"), FONT_TNR)
    rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


PARA1 = (
    "在线论坛的历史可追溯到个人计算机刚刚兴起的年代。"
    "1978 年 2 月，Ward Christensen 和 Randy Suess 在芝加哥搭建了 CBBS"
    "\uff08Computerized Bulletin Board System\uff09，"
    "这是有据可查的第一个拨号电子公告板系统，"
    "用户通过调制解调器拨号连接，以纯文本方式发布和阅读消息。"
    "此后一年，Tom Truscott 和 Jim Ellis 在杜克大学设计了 Usenet，于 1980 年正式上线。"
    "Usenet 采用分布式架构，消息在多台服务器之间自动转发，"
    "按 comp、sci、rec 等主题类别组织为新闻组，"
    "开创了以主题分区聚合讨论的模式。"
    "1984 年 Tom Jennings 创建的 FidoNet 进一步将独立运行的各 BBS 站点联结成网络，"
    "高峰期全球节点接近四万个。"
    "同一时期，CompuServe、GEnie 等商业在线服务以按小时计费的方式向公众开放论坛区，"
    "CompuServe 在 1990 年代中期拥有约 300 万付费订户，"
    "论坛由志愿者管理，覆盖技术、爱好和职业等多个领域。"
    "到 1994 年，仅美国便有约 6 万个 BBS 系统服务于 1700 万用户。"
    "这一阶段的系统虽然界面简陋、接入成本高，"
    "但已经确立了论坛的核心形态：用户注册、主题分区、按时间序排列的帖子与回复。"
)

PARA2 = (
    "1994 至 1995 年间，Mosaic 和 Netscape 浏览器的普及"
    "以及廉价拨号上网的出现，"
    "使基于万维网的论坛迅速取代了传统 BBS。"
    "1996 年发布的 WWWBoard 和 UBB\uff08Ultimate Bulletin Board\uff09"
    "是最早的浏览器端论坛软件，"
    "管理员无须编写代码即可在网页上创建和配置社区。"
    "此后 vBulletin\uff082000 年\uff09成为付费论坛软件的代表，"
    "phpBB 作为免费开源替代随之流行，"
    "Invision Power Board\uff082002 年\uff09和 SMF 等产品也相继推出。"
    "这些软件提供了统一的版面索引、分页帖子列表、用户资料页和站内消息等功能模块，"
    "Web 论坛由此进入标准化阶段。"
    "在中国，Discuz! 于 2001 年发布并开源，"
    "凭借丰富的插件和皮肤系统，"
    "占据了国内超过八成的论坛市场份额[1]。"
)

PARA3 = (
    "中国互联网论坛的建设几乎与 Web 论坛的全球兴起同步展开。"
    "1994 年曙光站作为中国大陆第一个 BBS 上线，"
    "1995 年水木清华 BBS 在教育网开放，"
    "注册用户一度达到 30 万，最高同时在线 23674 人。"
    "1998 年西祠胡同首创\u201c自由开版、自主管理\u201d模式，"
    "一年内在线人数突破百万。"
    "1999 年天涯社区上线，聚集了大批写作者，"
    "\u300a明朝那些事儿\u300b等知名连载即发源于此。"
    "2003 年百度贴吧将搜索引擎与论坛结合，"
    "上线两年后日均发帖量已达 200 万。"
    "高校场景中，水木清华和北大未名 BBS 曾是学生信息交流的核心平台，"
    "但 2005 年教育部要求高校 BBS 限制校外访问后，"
    "两者的用户规模和影响力均大幅收缩，"
    "水木清华于 2012 年正式关站[2]。"
)

PARA4 = (
    "2009 年以后，微博、微信等社交平台和移动端应用对论坛形成了持续冲击。"
    "天涯社区在鼎盛期日访问量约 2000 万，"
    "但未能完成移动端转型，"
    "2023 年 4 月停止服务，2024 年 2 月正式宣告破产。"
    "传统论坛的衰退并不意味着\u201c以主题聚合讨论\u201d的需求消失："
    "Reddit\uff082005 年创立\uff09以用户投票排序取代时间序，"
    "Stack Overflow\uff082008 年上线\uff09以声望积分和标签体系重新组织知识问答，"
    "两者至今仍保持高活跃度。"
    "在技术层面，Discourse\uff082013 年发布\uff09"
    "将论坛软件从 PHP 时代推进到 Ruby on Rails 和 Ember.js 的前后端分离架构，"
    "提供实时通知、Markdown 编辑和移动端自适应等现代特性。"
    "高校领域也出现了采用 Spring Boot 与 Vue 前后端分离方案的校园论坛系统，"
    "功能完整性和开发效率较传统架构有明显提升[5]。"
)

PARA5 = (
    "微服务架构将单体应用拆解为多个可独立部署的服务单元，"
    "服务之间经由轻量级协议通信，业务模块得以解耦并按需伸缩。"
    "Spring Cloud 体系在服务注册与发现、API 网关路由、"
    "流量控制与熔断降级等环节已形成较为完整的技术栈，"
    "辅以 Docker 容器化部署，可实现服务的标准化交付与弹性伸缩[6]。"
    "高校信息化建设同样受益于此。"
    "有学者将 Spring Cloud 微服务架构引入协同人员管理平台，"
    "借助服务拆分和消息队列实现了业务解耦与数据准实时同步[8]\uff1b"
    "科研管理领域亦出现了覆盖需求调研到成果转化全流程的微服务工程案例[9]。"
    "安全层面的研究也在推进："
    "Aldea 等人系统梳理了微服务环境下认证机制面临的漏洞与防护策略，"
    "强调 JWT 令牌的短生命周期设计和最小权限原则是缩小攻击面的关键手段[7]。"
    "面向高并发场景，已有论坛系统引入 Redis 缓存与消息队列，"
    "并结合 Nginx 负载均衡和容器化部署来提升大规模访问下的稳定性[21]。"
    "质量保障同样受到关注\u2014\u2014"
    "合理的测试策略与持续集成流程能显著降低线上故障率并加快需求变更的响应[10]。"
    "总体而言，论坛系统的技术选型已从拨号 BBS 经由 PHP 单体论坛"
    "演进到前后端分离与微服务架构相结合的阶段，"
    "但在高校场景中，将微服务架构完整应用于论坛系统"
    "并配套可量化、可复现的测试与运维体系的工程实践仍然较少，"
    "这正是本文的研究切入点。"
)

ALL_PARAS = [PARA1, PARA2, PARA3, PARA4, PARA5]

doc = Document("docs/论文修改.docx")

p16 = doc.paragraphs[16]
p16_elem = p16._element

p16.clear()
run = p16.add_run(ALL_PARAS[0])
set_run_font(run, FONT_SONG, FONT_TNR, PT_XIAOSI)
p16.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p16.paragraph_format.first_line_indent = Cm(0.85)
p16.paragraph_format.line_spacing = 1.5

prev_elem = p16_elem
for text in ALL_PARAS[1:]:
    new_p = make_body_para(text)
    prev_elem.addnext(new_p)
    prev_elem = new_p

doc.save("docs/论文修改.docx")

total = sum(len(t) for t in ALL_PARAS)
print(f"1.2 expanded: {len(ALL_PARAS)} paragraphs, {total} chars total")
